import streamlit as st
import os
import io
import subprocess
from docx import Document
from pypdf import PdfReader, PdfWriter

# Page UI properties layout configuration
st.set_page_config(
    page_title="Dr. AIT BDA Report Generator", 
    page_icon="🎓", 
    layout="centered"
)

# Professional Institutional Slate Theme Styling
st.markdown(
    """
    <style>
    .stApp {
        background-color: #0f172a;
    }
    h1, h2, h3, p, span, label, .stMarkdown {
        color: #f8fafc !important;
    }
    .stTextInput input {
        background-color: rgba(30, 41, 59, 0.7) !important;
        color: #f8fafc !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
    }
    div.stButton > button:first-child {
        background-color: #1e3a8a !important;
        color: white !important;
        border-radius: 6px;
        width: 100%;
        border: none;
        padding: 10px;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.title("🎓 Dr. AIT BDA Lab Report Compiler")
st.markdown("Enter your student details below to instantly generate a fully compiled, personalized, print-ready laboratory manual featuring your official front coversheets.")

st.markdown("---")

front_template_path = "BDA FRONT PAGE1_merged.docx"
manual_pdf_path = "BAD_Manual_merged.pdf"

# Verify repository readiness state inside the sidebar interface
st.sidebar.header("📦 System Resource Status")
resources_ready = True

if os.path.exists(front_template_path):
    st.sidebar.success("✅ Front Page Word Template Loaded")
else:
    st.sidebar.error("❌ Missing: BDA FRONT PAGE1_merged.docx")
    resources_ready = False

if os.path.exists(manual_pdf_path):
    st.sidebar.success("✅ Master Lab Manual Loaded")
else:
    st.sidebar.error("❌ Missing: BAD_Manual_merged.pdf")
    resources_ready = False

# --- STUDENT INPUT MATRIX ---
st.subheader("👤 Enter Student Credentials")
col1, col2 = st.columns(2)
student_name = col1.text_input("Full Student Name:", placeholder="e.g. CHETHAN PRASAD L")
student_usn = col2.text_input("University Seat Number (USN):", placeholder="e.g. 1DA25SCS04")

# Core function to scan paragraph formatting blocks and replace placeholders
def process_word_template(path, replacement_name, replacement_usn):
    doc = Document(path)
    
    # Original hardcoded placeholders located in your files
    search_usn = "1DA25SCS18"
    search_name_1 = "SRIPADA RAO H"
    search_name_2 = "H SRIPADA RAO"
    
    # Loop over standard paragraph runs
    for para in doc.paragraphs:
        if search_usn in para.text:
            for run in para.runs:
                if search_usn in run.text:
                    run.text = run.text.replace(search_usn, replacement_usn.upper())
        if search_name_1 in para.text or search_name_2 in para.text:
            for run in para.runs:
                if search_name_1 in run.text:
                    run.text = run.text.replace(search_name_1, replacement_name.upper())
                if search_name_2 in run.text:
                    run.text = run.text.replace(search_name_2, replacement_name.upper())

    # Loop over hidden cells inside tables (signature validation matrices)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if search_usn in para.text:
                        for run in para.runs:
                            if search_usn in run.text:
                                run.text = run.text.replace(search_usn, replacement_usn.upper())
                    if search_name_1 in para.text or search_name_2 in para.text:
                        for run in para.runs:
                            if search_name_1 in run.text:
                                run.text = run.text.replace(search_name_1, replacement_name.upper())
                            if search_name_2 in run.text:
                                run.text = run.text.replace(search_name_2, replacement_name.upper())
                                
    return doc

# --- GENERATION PIPELINE ENGINE ---
if resources_ready:
    if student_name.strip() and student_usn.strip():
        st.markdown("---")
        
        if st.button("⚡ Compile & Generate Full Printable Report"):
            with st.spinner("Processing templates, generating customized cover sheets, and stitching manual pages..."):
                
                # Define temporary working files
                temp_docx = "temp_modified_front.docx"
                temp_pdf = "temp_modified_front.pdf"
                final_output = f"{student_usn.upper()}_BDA_Complete_Report.pdf"
                
                try:
                    # 1. Update text metadata directly within the .docx file layers
                    modified_doc = process_word_template(front_template_path, student_name, student_usn)
                    modified_doc.save(temp_docx)
                    
                    # 2. Cross-platform headless CLI invocation to convert docx to pdf flawlessly
                    try:
                        # For Linux/Streamlit environments hosting LibreOffice modules
                        subprocess.run([
                            'libreoffice', '--headless', '--convert-to', 'pdf', temp_docx
                        ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    except Exception:
                        # Fallback for alternative environments utilizing standard conversion scripts
                        from docx2pdf import convert
                        convert(temp_docx, temp_pdf)
                    
                    # 3. Stitch the generated front sheets with the core manual using pypdf
                    pdf_writer = PdfWriter()
                    
                    front_reader = PdfReader(temp_pdf)
                    for page in front_reader.pages:
                        pdf_writer.add_page(page)
                        
                    main_manual_reader = PdfReader(manual_pdf_path)
                    for page in main_manual_reader.pages:
                        pdf_writer.add_page(page)
                    
                    # Save out the complete consolidated file
                    with open(final_output, "wb") as out_f:
                        pdf_writer.write(out_f)
                    
                    st.success("✨ Complete laboratory report compiled successfully with original coversheets!")
                    
                    # 4. Provide the instant file download utility button
                    with open(final_output, "rb") as final_bytes:
                        st.download_button(
                            label="📥 Download Finished Lab Report (PDF)",
                            data=final_bytes.read(),
                            file_name=final_output,
                            mime="application/pdf"
                        )
                        
                    # Clean up workspace cache elements silently
                    for file in [temp_docx, temp_pdf, final_output]:
                        if os.path.exists(file):
                            os.remove(file)
                            
                except Exception as engine_err:
                    st.error("🚨 An issue occurred during report compilation processing loops.")
                    st.exception(engine_err)
    else:
        st.info("💡 Please type in the Student Name and USN above to activate the automated compiler.")
